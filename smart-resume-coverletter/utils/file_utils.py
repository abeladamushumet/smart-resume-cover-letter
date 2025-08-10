import os
import io
from pathlib import Path
from typing import Tuple

import pdfplumber
from docx import Document

EXPORTS_DIR = Path(__file__).resolve().parent.parent / "exports"
EXPORTS_DIR.mkdir(parents=True, exist_ok=True)


def read_bytesio_as_text(file_like) -> str:
    # Fallback for generic text files
    content = file_like.read()
    if isinstance(content, bytes):
        try:
            return content.decode("utf-8", errors="ignore")
        except Exception:
            return content.decode("latin-1", errors="ignore")
    return str(content)


def read_uploaded_file(uploaded_file) -> Tuple[str, str]:
    """Return (text, detected_type) from a Streamlit UploadedFile or file-like.
    detected_type in {"pdf","docx","txt","unknown"}
    """
    filename = getattr(uploaded_file, "name", "uploaded")
    suffix = Path(filename).suffix.lower()

    # Move to bytes buffer for libraries that need a file path-like object
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)
    bytes_buffer = io.BytesIO(file_bytes)

    if suffix == ".pdf":
        try:
            text_parts = []
            with pdfplumber.open(bytes_buffer) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
            return ("\n".join(text_parts).strip(), "pdf")
        except Exception:
            # Fallback to raw decode
            return (file_bytes.decode("utf-8", errors="ignore"), "pdf")

    if suffix == ".docx":
        try:
            doc = Document(bytes_buffer)
            text = "\n".join(p.text for p in doc.paragraphs)
            return (text.strip(), "docx")
        except Exception:
            return (file_bytes.decode("utf-8", errors="ignore"), "docx")

    # Treat everything else as text
    try:
        return (file_bytes.decode("utf-8", errors="ignore"), "txt")
    except Exception:
        return (file_bytes.decode("latin-1", errors="ignore"), "unknown")


def save_text_as_docx(text: str, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(str(output_path))
    return output_path


def next_export_filename(prefix: str, suffix: str = ".docx") -> Path:
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    counter = 1
    while True:
        candidate = EXPORTS_DIR / f"{prefix}_{counter}{suffix}"
        if not candidate.exists():
            return candidate
        counter += 1