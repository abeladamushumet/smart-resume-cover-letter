from docx import Document

def create_sample_resume(path = r"C:\Users\hp\Desktop\VS code\smart-resume-cover-letter\sample_data\sample_resume.docx"
):
    doc = Document()
    doc.add_heading('John Doe', level=0)
    doc.add_paragraph('123 Main Street, City, Country')
    doc.add_paragraph('john.doe@email.com | +123 456 7890 | linkedin.com/in/johndoe')
    doc.add_paragraph()

    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(
        'Motivated software engineer with 3 years of experience in full-stack development, '
        'specializing in building scalable web applications and APIs. Skilled in Python, '
        'JavaScript, and cloud services.'
    )

    doc.add_heading('Work Experience', level=1)
    doc.add_heading('Software Engineer', level=2)
    doc.add_paragraph('Tech Solutions Inc. — Jan 2021 to Present')
    doc.add_paragraph(
        '- Developed and maintained web applications using React and Django.\n'
        '- Implemented RESTful APIs consumed by mobile and web clients.\n'
        '- Collaborated with cross-functional teams in Agile environment.'
    )

    doc.add_heading('Junior Developer', level=2)
    doc.add_paragraph('WebWorks Ltd. — Jun 2019 to Dec 2020')
    doc.add_paragraph(
        '- Assisted in developing client websites using HTML, CSS, and JavaScript.\n'
        '- Performed code reviews and bug fixes.'
    )

    doc.add_heading('Education', level=1)
    doc.add_paragraph('Bachelor of Science in Computer Science')
    doc.add_paragraph('State University — Graduated 2019')

    doc.add_heading('Skills', level=1)
    doc.add_paragraph(
        '- Programming: Python, JavaScript, SQL\n'
        '- Frameworks: React, Django, Flask\n'
        '- Tools: Git, Docker, AWS\n'
        '- Soft skills: Teamwork, Problem-solving, Communication'
    )

    doc.save(path)
    print(f'Sample resume saved to {path}')

if __name__ == "__main__":
    create_sample_resume()
